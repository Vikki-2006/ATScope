import re
import os
from typing import Dict, Any, List, Optional
import pdfplumber
import pypdf
import docx

# Standard comprehensive list of technical and professional skills for matching
COMMON_SKILLS = [
    # Programming Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "golang", "rust", "php", "ruby", "swift", "kotlin", "r", "scala", "html", "css", "sql", "bash", "shell",
    # Web Frameworks & Libraries
    "fastapi", "flask", "django", "react", "react.js", "next.js", "vue", "vue.js", "angular", "node.js", "express", "express.js", "spring boot", "asp.net", "tailwind css", "bootstrap", "alpine.js", "jquery", "graphql", "rest api", "web sockets",
    # Databases & Caching
    "postgresql", "postgres", "mysql", "sqlite", "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb", "oracle", "mariadb", "firebase", "supabase", "neo4j",
    # DevOps, Cloud & Infrastructure
    "docker", "kubernetes", "k8s", "aws", "amazon web services", "azure", "google cloud", "gcp", "terraform", "ansible", "jenkins", "github actions", "gitlab ci", "ci/cd", "nginx", "apache", "linux", "system design", "microservices", "serverless",
    # Data Science, AI & Machine Learning
    "machine learning", "deep learning", "artificial intelligence", "data analysis", "pandas", "numpy", "scikit-learn", "sklearn", "tensorflow", "pytorch", "keras", "spacy", "nltk", "opencv", "computer vision", "nlp", "natural language processing", "llm", "transformers", "hugging face", "langchain", "prompt engineering", "matplotlib", "seaborn", "tableau", "power bi",
    # Software Engineering & Methodology
    "git", "github", "gitlab", "bitbucket", "jira", "agile", "scrum", "kanban", "unit testing", "pytest", "jest", "cypress", "selenium", "tdd", "clean architecture", "oop", "algorithms", "data structures",
    # Professional & Soft Skills
    "leadership", "project management", "teamwork", "problem solving", "critical thinking", "communication", "time management", "stakeholder management"
]

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception:
            # Fallback to pypdf
            try:
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            except Exception as e:
                print(f"PDF extraction error: {e}")
        return text.strip()

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"DOCX extraction error: {e}")
        return text.strip()

    @classmethod
    def parse_file(cls, file_path: str, file_type: str) -> Dict[str, Any]:
        ext = file_type.lower()
        if "pdf" in ext:
            raw_text = cls.extract_text_from_pdf(file_path)
        elif "docx" in ext or "document" in ext:
            raw_text = cls.extract_text_from_docx(file_path)
        else:
            raw_text = ""
            
        return cls.parse_text(raw_text)

    @classmethod
    def parse_text(cls, text: str) -> Dict[str, Any]:
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        
        email = cls._extract_email(text)
        phone = cls._extract_phone(text)
        linkedin = cls._extract_linkedin(text)
        github = cls._extract_github(text)
        name = cls._extract_name(lines, email)
        skills = cls._extract_skills(text)
        sections = cls._extract_sections(text)
        
        education = cls._parse_education_section(sections.get("education", ""))
        experience = cls._parse_experience_section(sections.get("experience", ""))
        projects = cls._parse_projects_section(sections.get("projects", ""))
        certifications = cls._parse_list_section(sections.get("certifications", ""))
        languages = cls._parse_list_section(sections.get("languages", ""))
        summary = sections.get("summary", "")

        return {
            "name": name,
            "email": email,
            "phone": phone,
            "linkedin": linkedin,
            "github": github,
            "summary": summary,
            "skills": skills,
            "education": education,
            "experience": experience,
            "projects": projects,
            "certifications": certifications,
            "languages": languages,
            "raw_text": text
        }

    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(pattern, text)
        return match.group(0) if match else None

    @staticmethod
    def _extract_phone(text: str) -> Optional[str]:
        pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        match = re.search(pattern, text)
        return match.group(0) if match else None

    @staticmethod
    def _extract_linkedin(text: str) -> Optional[str]:
        pattern = r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?'
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    @staticmethod
    def _extract_github(text: str) -> Optional[str]:
        pattern = r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+/?'
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    @staticmethod
    def _extract_name(lines: List[str], email: Optional[str]) -> str:
        if not lines:
            return "Candidate"
        # First non-header line that doesn't contain email or phone is usually name
        for line in lines[:5]:
            if "@" not in line and not re.search(r'\d{5,}', line) and len(line.split()) <= 4:
                # Clean up punctuation
                clean = re.sub(r'[^a-zA-Z\s]', '', line).strip()
                if clean and len(clean) > 2:
                    return clean.title()
        return "Candidate"

    @staticmethod
    def _extract_skills(text: str) -> List[str]:
        found_skills = set()
        text_lower = text.lower()
        
        for skill in COMMON_SKILLS:
            # Word boundary search for precise matching
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                # Standardize display formatting
                found_skills.add(skill.title() if len(skill) > 3 else skill.upper())
                
        return sorted(list(found_skills))

    @staticmethod
    def _extract_sections(text: str) -> Dict[str, str]:
        sections = {
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "certifications": "",
            "languages": ""
        }
        
        lines = text.split("\n")
        current_section = None
        section_content = {k: [] for k in sections}
        
        section_headers = {
            "summary": ["summary", "profile", "objective", "about me", "professional summary"],
            "experience": ["experience", "work experience", "employment history", "work history", "professional experience"],
            "education": ["education", "academic background", "qualifications"],
            "skills": ["skills", "technical skills", "technologies", "core competencies"],
            "projects": ["projects", "personal projects", "key projects"],
            "certifications": ["certifications", "licenses", "certificates", "training"],
            "languages": ["languages", "language proficiency"]
        }

        for line in lines:
            clean_line = line.strip().lower()
            clean_hdr = re.sub(r'[^a-z\s]', '', clean_line).strip()
            
            matched_sec = None
            for sec_key, keywords in section_headers.items():
                if clean_hdr in keywords or any(clean_hdr == kw for kw in keywords):
                    matched_sec = sec_key
                    break
            
            if matched_sec:
                current_section = matched_sec
            elif current_section:
                section_content[current_section].append(line)

        for k in sections:
            sections[k] = "\n".join(section_content[k]).strip()
            
        return sections

    @staticmethod
    def _parse_education_section(text: str) -> List[Dict[str, Any]]:
        if not text:
            return []
        items = []
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        degrees = ["B.S.", "B.A.", "B.Tech", "Bachelor", "M.S.", "M.A.", "M.Tech", "Master", "Ph.D.", "Diploma", "High School"]
        
        for line in lines:
            for d in degrees:
                if d.lower() in line.lower():
                    items.append({"degree": line, "institution": "", "year": ""})
                    break
        if not items and lines:
            items.append({"degree": lines[0], "institution": lines[1] if len(lines) > 1 else "", "year": ""})
        return items

    @staticmethod
    def _parse_experience_section(text: str) -> List[Dict[str, Any]]:
        if not text:
            return []
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        exp_list = []
        current_exp = None

        for line in lines:
            # Check if line looks like a job title or company name
            if any(role in line.lower() for role in ["engineer", "developer", "manager", "intern", "lead", "architect", "analyst", "consultant", "specialist"]):
                if current_exp:
                    exp_list.append(current_exp)
                current_exp = {"title": line, "company": "", "duration": "", "description": []}
            elif current_exp:
                current_exp["description"].append(line)

        if current_exp:
            exp_list.append(current_exp)
            
        if not exp_list and lines:
            exp_list.append({"title": lines[0], "company": "", "duration": "", "description": lines[1:]})

        return exp_list

    @staticmethod
    def _parse_projects_section(text: str) -> List[Dict[str, Any]]:
        if not text:
            return []
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        projects = []
        for line in lines[:5]:  # Capture primary project titles
            if len(line) > 3:
                projects.append({"name": line, "description": line})
        return projects

    @staticmethod
    def _parse_list_section(text: str) -> List[str]:
        if not text:
            return []
        return [line.strip("-•* ") for line in text.split("\n") if line.strip()]
